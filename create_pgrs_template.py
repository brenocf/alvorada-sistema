
from docx import Document
import os

TEMPLATES_DIR = "templates"

def create_pgrs():
    doc = Document()
    doc.add_heading('PLANO DE GERENCIAMENTO DE RESÍDUOS SÓLIDOS (PGRS)', 0)
    
    doc.add_paragraph('EMPRESA: {RAZAO}')
    doc.add_paragraph('CNPJ: {CNPJ}')
    doc.add_paragraph('DATA: {DATA_HOJE}')
    
    doc.add_heading('1. INTRODUÇÃO', level=1)
    doc.add_paragraph('O presente plano visa estabelecer diretrizes para o manejo adequado dos resíduos gerados na atividade.')
    
    doc.add_heading('2. INVENTÁRIO DE RESÍDUOS', level=1)
    doc.add_paragraph('Abaixo listamos os resíduos gerados e suas respectivas destinações:')
    
    # Placeholder for the dynamic table/list
    doc.add_paragraph('{TABELA_RESIDUOS}')
    
    doc.add_heading('3. RESPONSABILIDADE', level=1)
    doc.add_paragraph('A empresa se compromete a seguir as normas da Lei 12.305/2010.')
    
    doc.add_paragraph('\n\n____________________________________________________')
    doc.add_paragraph('Responsável Técnico')
    
    path = os.path.join(TEMPLATES_DIR, 'Modelo_PGRS.docx')
    doc.save(path)
    print(f"Template criado em: {path}")

if __name__ == "__main__":
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
    create_pgrs()
