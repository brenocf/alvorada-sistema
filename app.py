import streamlit as st
import pandas as pd
import time
import importlib

# Módulos do Sistema
import database
import search_engine
import business_logic
import cnae_mapping

# Reload para desenvolvimento (opcional)
importlib.reload(database)
importlib.reload(business_logic)
importlib.reload(search_engine)
importlib.reload(search_engine)
importlib.reload(cnae_mapping)

# Imports para Documentos
import io
import os
from docx import Document
from datetime import datetime

# --- 1. CONFIGURAÇÃO INICIAL ---
# --- 1. CONFIGURAÇÃO INICIAL ---
st.set_page_config(
    page_title="Alvorada Projetos",
    page_icon="☀️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Inicializar Banco de Dados
database.init_db()

# --- 1.1 ESTADO DE SESSÃO & AUTH ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'user_name' not in st.session_state:
    st.session_state['user_name'] = ""

def login_page():
    c1, c2, c3 = st.columns([1, 1, 1])
    with c2:
        try:
            st.image("assets/logo.png", use_container_width=True)
        except:
            st.title("ALVORADA ☀️")
        
        tab_login, tab_cad = st.tabs(["Acessar", "Criar Conta"])
        
        with tab_login:
            with st.form("login_form"):
                u = st.text_input("Usuário")
                p = st.text_input("Senha", type="password")
                if st.form_submit_button("ENTRAR", use_container_width=True):
                    user_data = database.verify_login(u, p)
                    if user_data:
                        st.session_state['logged_in'] = True
                        st.session_state['user_name'] = user_data['name']
                        st.toast(f"Bem-vindo, {user_data['name']}!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("Usuário ou senha inválidos.")
        
        with tab_cad:
            st.warning("Área restrita a colaboradores.")
            with st.form("cad_form"):
                new_u = st.text_input("Novo Usuário")
                new_n = st.text_input("Nome Completo")
                new_p = st.text_input("Escolha uma Senha", type="password")
                if st.form_submit_button("REGISTRAR"):
                    if database.create_user(new_u, new_p, new_n):
                        st.success("Conta criada! Faça login.")
                    else:
                        st.error("Usuário já existe.")

if not st.session_state['logged_in']:
    login_page()
    st.stop() # Interrompe o script aqui se não estiver logado

# Custom CSS para Badges e Identidade ALVORADA
st.markdown("""
<style>
    /* PALETA ALVORADA: Navy #0B2545 | Orange #FF9F1C | White #FFFFFF */
    
    /* Fundo App */
    .stApp {
        background-color: #f8f9fa;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #0B2545;
    }
    [data-testid="stSidebar"] * {
        color: white !important;
    }
    
    /* Botões */
    div.stButton > button {
        background-color: #FF9F1C;
        color: #0B2545;
        border: none;
        font-weight: bold;
    }
    div.stButton > button:hover {
        background-color: #ffb75e;
        color: #0B2545;
    }

    /* Títulos */
    h1, h2, h3 { 
        color: #0B2545 !important; 
        font-family: 'Segoe UI', sans-serif;
    }
    
    /* Badges */
    .badge-risk { background-color: #ffcccc; color: #cc0000; padding: 4px 8px; border-radius: 4px; font-weight: bold; }
    .badge-ok { background-color: #ccffcc; color: #006600; padding: 4px 8px; border-radius: 4px; font-weight: bold; }
    .badge-neutral { background-color: #f0f0f0; color: #333; padding: 4px 8px; border-radius: 4px; }
</style>
""", unsafe_allow_html=True)
    

# --- 2. NAVEGAÇÃO & ESTILO PROFISSIONAL ---
# (Estilos já definidos no bloco anterior)

with st.sidebar:
    try:
        st.image("assets/logo.png", use_container_width=True)
    except:
        st.title("ALVORADA PROJETOS")
        
    st.markdown(f"**Olá, {st.session_state.get('user_name', 'Colaborador')}**")
    if st.button("Sair", key="logout_btn"):
        st.session_state['logged_in'] = False
        st.rerun()
        
    st.divider()
    
    menu = st.radio("MENU PRINCIPAL", 
             ["Dashboard", "Buscar Novos (API)", "Minha Carteira", "Financeiro & Propostas", "Fábrica de Documentos", "Calendário de Obrigações", "GED - Arquivos", "Monitor de Compliance", "Gerador PGRS"],
             captions=["Visão Geral e Métricas", "Importação de Leads", "Gestão de Clientes", "Orçamentos e Caixa", "Automação Doc", "Agenda Técnica", "Gestão de Documentos", "Vencimentos", "Plano de Resíduos"])
             
    st.divider()
    
    # Mini Dashboard na Sidebar
    try:
        df_side = database.get_carteira()
        total_side = len(df_side)
        novos_side = len(df_side[df_side['status_crm'] == 'Novo'])
        
        st.markdown("### Status Rápido")
        c1, c2 = st.columns(2)
        c1.metric("Total", total_side)
        c2.metric("Novos", novos_side)
        
        st.progress(min(total_side / 100, 1.0), text="Meta Mensal")
    except:
        pass


# --- PÁGINA: BUSCAR NOVOS ---
if menu == "Buscar Novos (API)":
    st.header("Buscar Novos Leads")
    st.info("Conecte-se à API ou use Mock para encontrar empresas e adicionar à sua carteira.")
    
    col_mode, col_key = st.columns([1, 2])
    modo = col_mode.selectbox("Fonte de Dados", ["Mock / Teste", "API CNPJá (Real)"])
    api_key = ""
    if modo == "API CNPJá (Real)":
        api_key = col_key.text_input("API Key", type="password", value="e32611ea-918b-4ed0-8b0e-e4432f9de77b-1df5f4c3-cd44-4821-ac4d-ce5b9c34271a")

    if st.button("Executar Varredura", type="primary"):
        with st.spinner("Buscando dados..."):
            raw_leads = []
            if modo == "Mock / Teste":
                raw_leads = search_engine.buscar_empresas(mock_mode=True)
            else:
                raw_leads, erro = search_engine.buscar_cnpja_comercial(api_key=api_key)
                if erro: st.error(erro)

            if raw_leads:
                # Processar Leads (Enriquecer)
                processed_leads = business_logic.analisar_leads(raw_leads)
                
                # Gerar Link de Rota na Memória
                for p in processed_leads:
                     if p.get('logradouro'):
                         log = str(p.get('logradouro','')).replace(' ','+')
                         num = str(p.get('numero','')).replace(' ','+')
                         bai = str(p.get('bairro','')).replace(' ','+')
                         p['Rota'] = f"https://www.google.com/maps/search/?api=1&query={log}+{num}+{bai}+Iguatu+CE"
                     else:
                         p['Rota'] = ""

                st.session_state['novos_leads'] = processed_leads
                st.success(f"{len(processed_leads)} empresas encontradas!")
    
    # Exibir Resultados Recentes da Memória
    if 'novos_leads' in st.session_state:
        # DEBUG INSPECTOR
        with st.expander("Inspecionar Resposta Real da API (Debug)"):
             if st.session_state['novos_leads']:
                 st.write("Dados brutos do primeiro item:")
                 st.json(st.session_state['novos_leads'][0])
                 
        df_novos = pd.DataFrame(st.session_state['novos_leads'])
        
        st.divider()
        st.subheader("Selecione para Importar")
        
        df_novos['Importar'] = True
        
        edited_df = st.data_editor(
            df_novos[["Importar", "razao_social", "grupo_descricao", "bairro", "tag_risco"]],
            use_container_width=True,
            hide_index=True
        )
        
        if st.button("Salvar Marcados na Carteira", type="primary"):
            stats = {"inserted": 0, "updated": 0, "skipped": 0, "error": 0}
            indices_selecionados = edited_df[edited_df['Importar'] == True].index
            
            prog_bar = st.progress(0)
            total = len(indices_selecionados)
            
            for i, idx in enumerate(indices_selecionados):
                lead_obj = st.session_state['novos_leads'][idx]
                cnpj = lead_obj['cnpj']
                
                # ENRIQUECIMENTO DE DADOS
                if not lead_obj.get('cnaes_secundarios'):
                    st.toast(f"Enriquecendo dados para {cnpj}...")
                    api_key = st.session_state.get('api_key_cnpja')
                    if api_key:
                        detalhes = search_engine.consultar_detalhes_cnpj(api_key, cnpj)
                        if detalhes:
                            count_sec = len(detalhes.get('cnaes_secundarios', []))
                            if count_sec > 0:
                                st.toast(f"Enriquecido: +{count_sec} ativ. secundárias.")
                            else:
                                st.toast(f"Detalhes obtidos, mas sem secundárias.")
                                
                            lead_obj.update(detalhes)
                            re_analysis = business_logic.analisar_leads([lead_obj])
                            if re_analysis:
                                lead_obj = re_analysis[0]
                                
                # Debug Temporário
                p_debug = lead_obj.get('porte_receita', 'KEY_MISSING')
                # st.toast(f"P. Rec: {p_debug}", icon="🐞")
                
                result = database.upsert_empresa(lead_obj)
                stats[result] = stats.get(result, 0) + 1
                prog_bar.progress((i + 1) / total)
            
            prog_bar.empty()
            
            msg = f"Processado com Sucesso. Novos: {stats['inserted']} | Atualizados: {stats['updated']} | Ignorados: {stats['skipped']}"
            st.toast(msg)
            time.sleep(2)
            st.rerun()


# --- PÁGINA: MINHA CARTEIRA ---
elif menu == "Minha Carteira":
    st.header("Minha Carteira")
    
    # Filtros
    col_f1, col_f2 = st.columns(2)
    bairro_filter = col_f1.text_input("Filtrar por Bairro:", placeholder="Ex: Centro")
    status_filter = col_f2.selectbox("Filtrar por Status", ["Todos", "Novo", "Em Negociação", "Cliente", "Descartado"])
    
    # Carregar Dados
    df_carteira = database.get_carteira()
    
    if bairro_filter:
        df_carteira = df_carteira[df_carteira['bairro'].str.contains(bairro_filter, case=False, na=False)]
    if status_filter != "Todos":
        df_carteira = df_carteira[df_carteira['status_crm'] == status_filter]
    
    if df_carteira.empty:
        st.warning("Nenhuma empresa encontrada com estes filtros.")
    else:
        # Layout Mestre-Detalhe
        col_list, col_detail = st.columns([1, 2.5])
        
        with col_list:
            st.markdown("### Empresas")
            
            def format_option(row):
                 return f"{row['nome_fantasia'] or row['razao_social']} | {row['bairro']}"
            
            options_map = {i: format_option(row) for i, row in df_carteira.iterrows()}
            
            selected_idx = st.radio("Selecione:", list(options_map.keys()), format_func=lambda x: options_map[x], label_visibility="collapsed")
            selected_cnpj = df_carteira.loc[selected_idx, 'cnpj']

        # --- VISÃO DE DETALHE (PRONTUÁRIO) ---
        with col_detail:
            # Container estilizado
            with st.container(border=True):
                empresa = database.get_empresa(selected_cnpj)
                
                if empresa:
                    # Cabeçalho Rico
                    c_head1, c_head2 = st.columns([3, 1])
                    nome_display = empresa.get('nome_fantasia') or empresa.get('razao_social')
                    c_head1.markdown(f"## {nome_display}")
                    c_head1.caption(f"Razão Social: {empresa['razao_social']} | CNPJ: {empresa['cnpj']}")
                    
                    # Status Button
                    current_status = empresa.get('status_crm', 'Novo')
                    new_st = c_head2.selectbox("Status", ["Novo", "Em Negociação", "Cliente", "Descartado"], index=["Novo", "Em Negociação", "Cliente", "Descartado"].index(current_status))
                    if new_st != current_status:
                        database.update_status_crm(selected_cnpj, new_st)
                        st.toast("Status Atualizado!")
                        time.sleep(0.5)
                        st.rerun()

                    # Badges Row
                    st.markdown("---")
                    col_b1, col_b2, col_b3 = st.columns(3)
                    
                    cor_risco = "badge-risk" if "ALTO" in str(empresa['risco']) else "badge-neutral"
                    col_b1.markdown(f"**Risco Ambiental**<br><span class='{cor_risco}'>{empresa['risco']}</span>", unsafe_allow_html=True)
                    
                    is_isento = "ISENTO" in str(empresa['status_taxa'])
                    cor_taxa = "badge-ok" if is_isento else "badge-neutral"
                    col_b2.markdown(f"**Taxa de Licenciamento**<br><span class='{cor_taxa}'>{empresa['status_taxa']}</span>", unsafe_allow_html=True)
                    
                    data_abert = empresa.get('data_abertura', 'N/A')
                    col_b3.markdown(f"**Abertura**<br>{data_abert}", unsafe_allow_html=True)
                    
                    st.markdown("---")

                    # Tabs Profissionais
                    tab1, tab2, tab3 = st.tabs(["Dados e Contato", "Diagnóstico Técnico", "CRM e Histórico"])
                    
                    with tab1:
                        c1, c2 = st.columns(2)
                        c1.markdown("#### Endereço")
                        c1.write(f"{empresa['logradouro']}, {empresa['numero']}")
                        c1.write(f"{empresa['bairro']} - Iguatu/CE")
                        c1.write(f"CEP: {empresa['cep']}")
                        
                        if empresa.get('rota_link'):
                            c1.link_button("Abrir no Google Maps", empresa['rota_link'])
                            
                        c2.markdown("#### Contato")
                        tel = empresa.get('telefone') or 'Não informado'
                        c2.write(f"Tel: {tel}")
                        if tel != 'Não informado':
                             phone_clean = "".join(filter(str.isdigit, str(tel)))
                             c2.link_button("WhatsApp", f"https://wa.me/55{phone_clean}")
                        
                        st.markdown("#### Sócios (QSA)")
                        st.info(empresa.get('qsa') or "Não disponível")

                    with tab2:
                        import json
                        dados_extra = json.loads(empresa.get('dados_extra', '{}'))
                        
                        # CNAE Principal
                        st.info(f"**Atividade Principal:** {dados_extra.get('cnae_fiscal_principal')} - {dados_extra.get('cnae_fiscal_descricao')}")
                        
                        # Secundários
                        secundarios = dados_extra.get('cnaes_secundarios', [])
                        if secundarios:
                            st.write(f"**{len(secundarios)} Atividades Secundárias:**")
                            # Lógica de badge de risco aqui
                            for sec in secundarios:
                                st.code(sec, language="text")
                        else:
                            st.caption("Sem secundárias.")
                            
                        st.markdown("#### Dados para Cálculo")
                        
                        # Capital Social Formatado (R$)
                        cap_social = float(dados_extra.get('capital_social', 0))
                        cap_fmt = "R$ {:,.2f}".format(cap_social).replace(",", "X").replace(".", ",").replace("X", ".")
                        st.write(f"**Capital Social:** {cap_fmt}")
                        
                        st.write(f"**Grupo:** {empresa['grupo_atividade']}")
                        
                        # Exibir os dois Portes
                        col_p1, col_p2 = st.columns(2)
                        
                        p_amb = empresa['porte'] # Mapeado do porte_calculado
                        p_rec = dados_extra.get('porte_receita', 'Não Identificado')
                        
                        col_p1.info(f"**Porte Ambiental:**\n{p_amb}")
                        col_p2.write(f"**Porte Receita:**\n{p_rec}")

                    with tab3:
                        # Timeline CRM
                        st.markdown("#### Nova Interação")
                        with st.form("crm_log"):
                            col_i1, col_i2 = st.columns([1, 4])
                            tipo_int = col_i1.selectbox("Tipo", ["Nota", "Ligação", "Visita", "WhatsApp", "Email"])
                            txt_int = col_i2.text_input("Resumo", placeholder="Ex: Agendamos visita, enviamos proposta...")
                            if st.form_submit_button("Salvar Log"):
                                database.add_interacao(selected_cnpj, tipo_int, txt_int)
                                st.success("Salvo!")
                                time.sleep(0.5)
                                st.rerun()
                        
                        st.markdown("#### Histórico")
                        historico = database.get_historico(selected_cnpj)
                        if not historico.empty:
                            for _, row in historico.iterrows():
                                st.text(f"{row['data_hora']} | {row['tipo'].upper()} | {row['notas']}")
                        else:
                            st.caption("Nada registrado ainda.")


# --- PÁGINA: DASHBOARD ---
elif menu == "Dashboard":
    st.title("Painel de Controle")
    st.markdown("Visão geral da sua carteira de clientes.")
    
    df = database.get_carteira()
    
    if not df.empty:
        # Linha de Totais
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Total de Leads", len(df))
        m2.metric("Novos", len(df[df['status_crm'] == 'Novo']))
        m3.metric("Em Negociação", len(df[df['status_crm'] == 'Em Negociação']))
        m4.metric("Clientes Fechados", len(df[df['status_crm'] == 'Cliente']), delta_color="normal")
        
        st.divider()
        
        # Gráficos
        g1, g2 = st.columns(2)
        
        with g1:
            st.subheader("Funil de Vendas")
            status_counts = df['status_crm'].value_counts()
            st.bar_chart(status_counts, color="#4CAF50")
            
        with g2:
            st.subheader("Distribuição por Bairro")
            st.bar_chart(df['bairro'].value_counts().head(10))
            
        st.subheader("Últimos Adicionados")
        st.dataframe(
            df[['razao_social', 'bairro', 'risco', 'data_abertura']].head(5),
            use_container_width=True,
            hide_index=True
        )
            
    else:
        st.info("Ainda não há dados na carteira. Vá em 'Buscar Novos' para começar!")

# --- PÁGINA: FINANCEIRO & PROPOSTAS ---
elif menu == "Financeiro & Propostas":
    st.title("Gestão Financeira & Propostas")
    st.markdown("Gerador de orçamentos e controle de caixa.")
    
    # 1. Métricas do Topo
    df_prop = database.get_propostas()
    
    total_proposto = df_prop['valor'].sum() if not df_prop.empty else 0.0
    total_pago = df_prop[df_prop['status'] == 'Pago']['valor'].sum() if not df_prop.empty else 0.0
    a_receber = df_prop[df_prop['status'] == 'Aberto']['valor'].sum() if not df_prop.empty else 0.0
    
    c1, c2, c3 = st.columns(3)
    c1.metric("Propostas Emitidas", f"R$ {total_proposto:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
    c2.metric("Caixa Realizado (Pago)", f"R$ {total_pago:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."), delta="Entradas")
    c3.metric("A Receber (Aberto)", f"R$ {a_receber:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."), delta="Previsto")
    
    st.divider()
    
    col_nova, col_lista = st.columns([1, 2])
    
    # 2. Gerador de Orçamento
    with col_nova:
        st.subheader("Nova Proposta")
        with st.form("form_proposta"):
            # Carregar Clientes para Selectbox
            df_clients = database.get_carteira()
            lista_clientes = []
            if not df_clients.empty:
                lista_clientes = df_clients['razao_social'].tolist()
            
            cliente_sel = st.selectbox("Cliente", lista_clientes if lista_clientes else ["Nenhum cliente cadastrado"])
            
            # Tabela de Preços Sugerida
            TABELA_PRECOS = {
                "Licença Simplificada (LP/LI/LO)": 1500.00,
                "Renovação de Licença": 800.00,
                "Regularização (RAMA Anual)": 500.00,
                "Defesa de Auto de Infração": 2500.00,
                "Consultoria Mensal (Recorrente)": 600.00,
                "Personalizado / Outro": 0.00
            }
            
            servico_sel = st.selectbox("Serviço", list(TABELA_PRECOS.keys()))
            valor_sugerido = TABELA_PRECOS[servico_sel]
            
            valor_final = st.number_input("Valor (R$)", value=float(valor_sugerido), step=50.0)
            vencimento = st.date_input("Vencimento")
            
            submitted = st.form_submit_button("Gerar Orçamento 📄")
            if submitted and cliente_sel != "Nenhum cliente cadastrado":
                # Buscar CNPJ do cliente selecionado
                cnpj_target = df_clients[df_clients['razao_social'] == cliente_sel].iloc[0]['cnpj']
                
                database.add_proposta(
                    cnpj=cnpj_target,
                    cliente_nome=cliente_sel,
                    servico=servico_sel,
                    valor=valor_final,
                    data_vencimento=str(vencimento),
                    status="Aberto"
                )
                st.success("Proposta registrada!")
                time.sleep(1)
                st.rerun()
                
    # 3. Lista de Propostas
    with col_lista:
        st.subheader("Últimas Propostas")
        if not df_prop.empty:
            # Exibição Simplificada
            for _, row in df_prop.iterrows():
                with st.expander(f"{row['data_criacao'][:10]} | {row['cliente_nome']} - {row['servico']}"):
                    kc1, kc2, kc3 = st.columns(3)
                    val_fmt = f"R$ {row['valor']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                    kc1.metric("Valor", val_fmt)
                    kc2.write(f"**Vence:** {row['data_vencimento']}")
                    
                    # Controle de Status
                    current_stats = row['status']
                    new_stat = kc3.selectbox("Status", ["Aberto", "Pago", "Atrasado", "Cancelado"], 
                                           key=f"st_{row['id']}", 
                                           index=["Aberto", "Pago", "Atrasado", "Cancelado"].index(current_stats))
                    
                    if new_stat != current_stats:
                        database.update_status_proposta(row['id'], new_stat)
                        st.toast("Status atualizado!")
                        time.sleep(0.5)
                        st.rerun()
        else:
            st.info("Nenhuma proposta emitida.")

# --- PÁGINA: FÁBRICA DE DOCUMENTOS ---
elif menu == "Fábrica de Documentos":
    st.title("Fábrica de Documentos 📄")
    st.markdown("Gere contratos, procurações e requerimentos automaticamente.")
    
    c_conf, c_preview = st.columns([1, 1])
    
    with c_conf:
        st.subheader("1. Configuração")
        
        # Seletor de Cliente
        df_clients = database.get_carteira()
        lista_clientes = []
        if not df_clients.empty:
            lista_clientes = df_clients['razao_social'].tolist()
            
        cliente_sel = st.selectbox("Selecione o Cliente", lista_clientes if lista_clientes else ["Sem clientes"])
        
        # Seletor de Template
        try:
            templates = [f for f in os.listdir("templates") if f.endswith(".docx")]
        except FileNotFoundError:
            templates = []
            st.error("Pasta 'templates' não encontrada.")
            
        template_sel = st.selectbox("Selecione o Modelo", templates if templates else ["Sem modelos"])
        
        # Inputs Extras
        st.markdown("#### Campos Variáveis")
        servico_input = st.text_input("Serviço (para contratos)", "Licenciamento Ambiental")
        valor_input = st.number_input("Valor (R$)", 1500.0, step=100.0)
        
        btn_gerar = st.button("Processar Documento ⚙️")

    with c_preview:
        st.subheader("2. Download")
        
        if btn_gerar and cliente_sel != "Sem clientes" and template_sel != "Sem modelos":
            try:
                # 1. Carregar Dados do Cliente
                dados_cli = df_clients[df_clients['razao_social'] == cliente_sel].iloc[0]
                
                # 2. Preparar Dicionário de Substituição
                mapa_dados = {
                    "{RAZAO_SOCIAL}": str(dados_cli['razao_social']),
                    "{CNPJ}": str(dados_cli['cnpj']),
                    "{ENDERECO}": f"{dados_cli.get('logradouro','')} {dados_cli.get('numero','')} - {dados_cli.get('bairro','')}",
                    "{SERVICO}": servico_input,
                    "{VALOR}": f"{valor_input:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
                    "{DATA_HOJE}": datetime.now().strftime("%d/%m/%Y"),
                    "{ATIVIDADE}": str(dados_cli.get('descricao_atividade', 'Atividade Comercial'))
                }
                
                # 3. Processar DOCX
                doc_path = os.path.join("templates", template_sel)
                doc = Document(doc_path)
                
                def replace_paragraph(paragraph):
                    for key, val in mapa_dados.items():
                        if key in paragraph.text:
                            # Substituição simples (perde formatação se complexo, mas funciona p/ texto corrido)
                            paragraph.text = paragraph.text.replace(key, val)
                            
                # Iterar parágrafos
                for p in doc.paragraphs:
                    replace_paragraph(p)
                    
                # Iterar tabelas (se houver)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_paragraph(p)
                
                # 4. Salvar em Memória
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.success("Documento gerado com sucesso!")
                
                file_name = f"{template_sel.replace('.docx','')}_{cliente_sel[:10].strip()}.docx"
                
                st.download_button(
                    label="⬇️ Baixar Arquivo Editável (.docx)",
                    data=buffer,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Erro ao gerar: {e}")
        elif btn_gerar:
            st.warning("Selecione um cliente e um modelo válido.")

# --- PÁGINA: CALENDÁRIO DE OBRIGAÇÕES ---
elif menu == "Calendário de Obrigações":
    st.title("Agenda de Obrigações 📅")
    st.markdown("Controle de prazos, renovações e relatórios.")
    
    col_tools, col_cal = st.columns([1, 2])
    
    with col_tools:
        st.subheader("Ferramentas")
        
        # Gerador Automático
        st.markdown("##### 🚀 Automação")
        if st.button("Escanear Carteira & Gerar Tarefas"):
            carteira = database.get_carteira()
            novas = 0
            for _, emp in carteira.iterrows():
                # Regra 1: RAMA (Anual) para empresas > 1 ano
                abertura = str(emp.get('data_abertura', ''))
                if abertura and len(abertura) >= 10:
                    try:
                        dt_abertura = datetime.strptime(abertura[:10], "%Y-%m-%d")
                        anos = (datetime.now() - dt_abertura).days / 365
                        if anos >= 1:
                            # Adicionar tarefa RAMA para 31/03 do proximo ano
                            # Simplificado: 30 dias a partir de hoje
                            prazo = datetime.now() + pd.Timedelta(days=30)
                            database.add_obrigacao(emp['cnpj'], "Entrega do RAMA", prazo.strftime("%Y-%m-%d"), "Relatório")
                            novas += 1
                            
                        # Regra 2: Monitoramento Semestral (Exemplo)
                        database.add_obrigacao(emp['cnpj'], "Visita de Monitoramento", (datetime.now() + pd.Timedelta(days=15)).strftime("%Y-%m-%d"), "Visita")
                        novas += 1
                    except:
                        pass
            st.success(f"{novas} tarefas geradas!")
            time.sleep(1)
            st.rerun()
            
        st.divider()
        
        st.markdown("##### ➕ Nova Tarefa Manual")
        with st.form("nova_tarefa"):
            # Select Cliente
            df_cli = database.get_carteira()
            nomes = df_cli['razao_social'].tolist() if not df_cli.empty else []
            cli_task = st.selectbox("Cliente", nomes)
            titulo_task = st.text_input("Título", "Renovação de Licença")
            tipo_task = st.selectbox("Tipo", ["Licença", "Visita", "Relatório", "Taxa", "Outro"])
            prazo_task = st.date_input("Prazo Final")
            
            if st.form_submit_button("Agendar"):
                if cli_task:
                    cnpj = df_cli[df_cli['razao_social'] == cli_task].iloc[0]['cnpj']
                    database.add_obrigacao(cnpj, titulo_task, str(prazo_task), tipo_task)
                    st.success("Agendado!")
                    time.sleep(0.5)
                    st.rerun()
    
    with col_cal:
        st.subheader("Próximos Vencimentos")
        
        # Filtros
        mostrar_conc = st.checkbox("Mostrar Concluídos", value=False)
        
        df_tasks = database.get_obrigacoes(filtro_concluido=mostrar_conc)
        
        if not df_tasks.empty:
            for _, row in df_tasks.iterrows():
                # Definir Cor do Card
                prazo = datetime.strptime(row['data_prazo'], "%Y-%m-%d")
                hoje = datetime.now()
                delta = (prazo - hoje).days
                
                cor_borda = "#ccc" # default
                txt_status = "No Prazo"
                
                if row['concluido'] == 1:
                    cor_borda = "#4CAF50" # Green
                    txt_status = "Concluído"
                elif delta < 0:
                    cor_borda = "#F44336" # Red
                    txt_status = f"ATRASADO ({abs(delta)} dias)"
                elif delta <= 7:
                    cor_borda = "#FF9800" # Orange
                    txt_status = "Vence em breve"
                else: 
                    cor_borda = "#2196F3" # Blue
                    
                # Renderizar Card Customizado
                st.markdown(f"""
                <div style="border-left: 5px solid {cor_borda}; padding: 10px; background-color: #f9f9f9; margin-bottom: 10px; border-radius: 4px;">
                    <strong>{row['data_prazo']}</strong> | {row['razao_social']} <br>
                    <span style="font-size: 1.1em;">{row['titulo']}</span> <span style="font-size: 0.8em; color: gray;">({row['tipo']})</span><br>
                    <small>Status: {txt_status}</small>
                </div>
                """, unsafe_allow_html=True)
                
                if row['concluido'] == 0:
                    if st.button("Concluir ✅", key=f"done_{row['id']}"):
                        database.concluir_obrigacao(row['id'])
                        st.rerun()
        else:
            st.info("Nenhuma obrigação pendente.")

# --- PÁGINA: GED (ARQUIVO DIGITAL) ---
elif menu == "GED - Arquivos":
    st.title("GED - Gestão Eletrônica de Documentos 📂")
    st.markdown("Armazenamento seguro de licenças, protocolos e contratos.")
    
    col_upload, col_view = st.columns([1, 2])
    
    # Seletor de Cliente Global para a Tela
    df_cli = database.get_carteira()
    nomes = df_cli['razao_social'].tolist() if not df_cli.empty else []
    
    with col_upload:
        st.subheader("Upload de Arquivo")
        cli_sel = st.selectbox("Selecione a Empresa", nomes, key="ged_cli_sel")
        
        uploaded_file = st.file_uploader("Escolha um arquivo (PDF, DOCX, JPG)", type=["pdf", "docx", "jpg", "png"])
        tipo_arq = st.selectbox("Tipo de Documento", ["Licença Emitida", "Protocolo", "Contrato Assinado", "Planta/Mapa", "Outro"])
        
        if st.button("Salvar Arquivo 💾"):
            if cli_sel and uploaded_file:
                # 1. Definir Caminho
                cnpj_target = df_cli[df_cli['razao_social'] == cli_sel].iloc[0]['cnpj']
                safe_cnpj = cnpj_target.replace(".", "").replace("/", "").replace("-", "")
                
                # Criar pasta da empresa se não existir
                empresa_dir = os.path.join("docs", safe_cnpj)
                if not os.path.exists(empresa_dir):
                    os.makedirs(empresa_dir)
                    
                # 2. Salvar Físico
                file_path = os.path.join(empresa_dir, uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                    
                # 3. Salvar Banco
                database.add_arquivo(cnpj_target, uploaded_file.name, file_path, tipo_arq)
                
                st.success("Arquivo arquivado com sucesso!")
                time.sleep(1)
                st.rerun()
            else:
                st.warning("Selecione empresa e arquivo.")
                
    with col_view:
        st.subheader("Arquivos da Empresa")
        if cli_sel:
            cnpj_target = df_cli[df_cli['razao_social'] == cli_sel].iloc[0]['cnpj']
            df_files = database.get_arquivos(cnpj_target)
            
            if not df_files.empty:
                for _, row in df_files.iterrows():
                    with st.expander(f"{row['data_upload'][:10]} | {row['nome_arquivo']} ({row['tipo']})"):
                        st.write(f"Caminho: `{row['caminho_fisico']}`")
                        
                        # Botão de Download (Lê do disco)
                        try:
                            with open(row['caminho_fisico'], "rb") as f:
                                st.download_button(
                                    label="Baixar Arquivo ⬇️",
                                    data=f,
                                    file_name=row['nome_arquivo']
                                )
                        except FileNotFoundError:
                            st.error("Arquivo não encontrado no disco (pode ter sido movido).")
            else:
                st.info("Nenhum arquivo digitalizado para esta empresa.")

# --- PÁGINA: MONITOR DE COMPLIANCE ---
elif menu == "Monitor de Compliance":
    st.title("Monitor de Condicionantes & Prazos ⚖️")
    st.markdown("Painel de Controle de Vencimentos e Notificações.")
    
    # Métricas
    df_obs = database.get_obrigacoes()
    if not df_obs.empty:
        # Converter prazo para datetime
        df_obs['data_prazo'] = pd.to_datetime(df_obs['data_prazo'])
        hoje = pd.to_datetime(datetime.now().date())
        
        # Criar coluna de Dias Restantes
        df_obs['dias_restantes'] = (df_obs['data_prazo'] - hoje).dt.days
        
        vencidas = df_obs[df_obs['dias_restantes'] < 0]
        urgentes = df_obs[(df_obs['dias_restantes'] >= 0) & (df_obs['dias_restantes'] <= 30)]
        no_prazo = df_obs[df_obs['dias_restantes'] > 30]
        
        m1, m2, m3 = st.columns(3)
        m1.metric("VENCIDAS 🚨", len(vencidas), delta_color="inverse")
        m2.metric("Atenção (30 dias) ⚠️", len(urgentes), delta_color="off")
        m3.metric("No Prazo ✅", len(no_prazo))
        
        st.divider()
        
        st.subheader("Lista de Pendências")
        
        filtro_view = st.radio("Filtrar por Status", ["Todas", "Vencidas", "Urgentes"], horizontal=True)
        
        df_show = df_obs
        if filtro_view == "Vencidas":
            df_show = vencidas
        elif filtro_view == "Urgentes":
            df_show = urgentes
            
        if not df_show.empty:
            for _, row in df_show.iterrows():
                # Card de Obrigação
                dias = row['dias_restantes']
                cor = "#F44336" if dias < 0 else ("#FF9800" if dias <= 30 else "#4CAF50")
                msg_status = "VENCIDO" if dias < 0 else f"Vence em {dias} dias"
                
                with st.container():
                     c_info, c_action = st.columns([3, 1])
                     with c_info:
                         st.markdown(f"""
                         <div style="border-left: 5px solid {cor}; padding: 10px; background-color: white; border-radius: 4px; box-shadow: 0 1px 2px rgba(0,0,0,0.1);">
                             <strong>{row['razao_social']}</strong> | {row['tipo']}<br>
                             <span style="font-size: 1.2em; font-weight: bold;">{row['titulo']}</span><br>
                             <small>Prazo: {row['data_prazo'].strftime('%d/%m/%Y')} | Status: <span style="color:{cor}; font-weight:bold;">{msg_status}</span></small>
                         </div>
                         <br>
                         """, unsafe_allow_html=True)
                     
                     with c_action:
                         # Gerador de Link WhatsApp
                         phone_raw = database.get_empresa(row['cnpj_empresa']).get('telefone', '')
                         # Limpar telefone
                         phone = "".join([c for c in str(phone_raw) if c.isdigit()])
                         
                         msg_text = f"Olá, empresa {row['razao_social']}. Gostaríamos de lembrar sobre o vencimento de '{row['titulo']}' programado para {row['data_prazo'].strftime('%d/%m/%Y')}. Podemos agendar?"
                         import urllib.parse
                         msg_encoded = urllib.parse.quote(msg_text)
                         
                         link_wa = f"https://wa.me/55{phone}?text={msg_encoded}"
                         
                         st.markdown(f'''
                             <a href="{link_wa}" target="_blank">
                                 <button style="
                                     background-color: #25D366; 
                                     color: white; 
                                     border: none; 
                                     padding: 10px; 
                                     border-radius: 5px; 
                                     cursor: pointer; 
                                     width: 100%;
                                     font-weight: bold;">
                                     📱 Notificar
                                 </button>
                             </a>
                         ''', unsafe_allow_html=True)
        else:
            st.success("Nenhuma pendência neste filtro!")
            
    else:
        st.info("Nenhuma condcionante cadastrada no sistema.")

# --- PÁGINA: GERADOR DE PGRS (WIZARD) ---
elif menu == "Gerador PGRS":
    st.title("Gerador de PGRS ♻️")
    st.markdown("Criação de Planos de Gerenciamento de Resíduos.")
    
    # Wizard State
    if 'pgrs_step' not in st.session_state:
        st.session_state['pgrs_step'] = 1
    if 'pgrs_data' not in st.session_state:
        st.session_state['pgrs_data'] = {}
        
    # Barra de Progresso
    passos = ["1. Empresa", "2. Resíduos", "3. Destinação", "4. Conclusão"]
    st.progress(st.session_state['pgrs_step'] / 4, text=passos[st.session_state['pgrs_step']-1])
    
    # --- PASSO 1: EMPRESA ---
    if st.session_state['pgrs_step'] == 1:
        st.subheader("Passo 1: Selecionar Gerador")
        df_cli = database.get_carteira()
        nomes = df_cli['razao_social'].tolist() if not df_cli.empty else []
        
        cli = st.selectbox("Cliente", nomes)
        
        if st.button("Próximo ➡️"):
            if cli:
                st.session_state['pgrs_data']['cliente'] = cli
                st.session_state['pgrs_step'] = 2
                st.rerun()
            else:
                st.warning("Selecione um cliente.")
                
    # --- PASSO 2: RESÍDUOS ---
    elif st.session_state['pgrs_step'] == 2:
        st.subheader("Passo 2: Inventário de Resíduos")
        st.markdown("Selecione os resíduos gerados:")
        
        residuos = ["Papel/Papelão", "Plástico", "Vidro", "Metais", "Orgânico", "Resíduos Perigosos (Classe I)", "Lâmpadas", "Óleos/Graxas"]
        selecionados = []
        
        c1, c2 = st.columns(2)
        for i, r in enumerate(residuos):
            col = c1 if i < 4 else c2
            if col.checkbox(r, key=f"res_{i}"):
                selecionados.append(r)
                
        c_nav1, c_nav2 = st.columns([1,1])
        if c_nav1.button("⬅️ Voltar"):
            st.session_state['pgrs_step'] = 1
            st.rerun()
        if c_nav2.button("Próximo ➡️"):
            if selecionados:
                st.session_state['pgrs_data']['residuos'] = selecionados
                st.session_state['pgrs_step'] = 3
                st.rerun()
            else:
                st.warning("Selecione pelo menos um resíduo.")

    # --- PASSO 3: DESTINAÇÃO ---
    elif st.session_state['pgrs_step'] == 3:
        st.subheader("Passo 3: Coletores e Destinação")
        st.markdown("Informe quem coleta cada resíduo:")
        
        destinos = {}
        for res in st.session_state['pgrs_data']['residuos']:
            destinos[res] = st.text_input(f"Coletor para {res}", placeholder="Ex: Cooperativa ABC ou Prefeitura")
            
        c_nav1, c_nav2 = st.columns([1,1])
        if c_nav1.button("⬅️ Voltar"):
            st.session_state['pgrs_step'] = 2
            st.rerun()
            
        if c_nav2.button("GERAR PGRS 📝"):
            st.session_state['pgrs_data']['destinos'] = destinos
            st.session_state['pgrs_step'] = 4
            st.rerun()

    # --- PASSO 4: GERAÇÃO ---
    elif st.session_state['pgrs_step'] == 4:
        st.subheader("Passo 4: Relatório Gerado!")
        
        try:
            # Lógica de Geração
            data = st.session_state['pgrs_data']
            df_cli = database.get_carteira()
            info_cli = df_cli[df_cli['razao_social'] == data['cliente']].iloc[0]
            
            # Montar Texto da Tabela
            txt_residuos = ""
            for res, coletor in data['destinos'].items():
                txt_residuos += f"- {res}: Coletado por {coletor}\n"
                
            mapa = {
                "{RAZAO}": str(info_cli['razao_social']),
                "{CNPJ}": str(info_cli['cnpj']),
                "{DATA_HOJE}": datetime.now().strftime("%d/%m/%Y"),
                "{TABELA_RESIDUOS}": txt_residuos
            }
            
            # Carregar Template
            doc = Document("templates/Modelo_PGRS.docx")
            
             # Substituição Simples
            def replace_in_p(p):
                for k, v in mapa.items():
                    if k in p.text:
                        p.text = p.text.replace(k, v)
            
            for p in doc.paragraphs:
                replace_in_p(p)
                
            # Salvar
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.success("PGRS Criado com Sucesso!")
            st.download_button("⬇️ Baixar Plano de Resíduos", data=buffer, file_name=f"PGRS_{data['cliente'][:10].strip()}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            if st.button("Novo Plano"):
                st.session_state['pgrs_step'] = 1
                st.rerun()
                
        except Exception as e:
            st.error(f"Erro: {e}") 
            if st.button("Tentar Novamente"):
                 st.session_state['pgrs_step'] = 1
                 st.rerun()
