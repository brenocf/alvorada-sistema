
from docx import Document
from docx.shared import Pt
import os

TEMPLATES_DIR = "templates"

def create_contrato():
    doc = Document()
    doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS AMBIENTAIS', 0)
    
    p = doc.add_paragraph('Pelo presente instrumento, a empresa ')
    p.add_run('{RAZAO_SOCIAL}').bold = True
    p.add_run(', inscrita no CNPJ sob nº ')
    p.add_run('{CNPJ}').bold = True
    p.add_run(', estabelecida em {ENDERECO}, neste ato representada por seus sócios, doravante denominada CONTRATANTE.')
    
    doc.add_heading('1. DO OBJETO', level=1)
    doc.add_paragraph('O presente contrato tem por objeto a prestação de serviços de Consultoria Ambiental para fins de {SERVICO}.')
    
    doc.add_heading('2. DO VALOR', level=1)
    doc.add_paragraph('Pelos serviços prestados, a CONTRATANTE pagará o valor de R$ {VALOR}.')
    
    doc.add_paragraph('\n\n____________________________________________________')
    doc.add_paragraph('ASSINATURA DO REPRESENTANTE LEGAL')
    
    doc.save(os.path.join(TEMPLATES_DIR, 'Contrato_Servicos.docx'))

def create_procuracao():
    doc = Document()
    doc.add_heading('PROCURAÇÃO PARA FINS AMBIENTAIS', 0)
    
    p = doc.add_paragraph('OUTORGANTE: ')
    p.add_run('{RAZAO_SOCIAL}').bold = True
    p.add_run(', CNPJ: {CNPJ}, Endereço: {ENDERECO}.')
    
    doc.add_paragraph('Pelo presente instrumento particular de procuração, a Outorgante nomeia e constitui seu bastante procurador o Eng. Ambiental [SEU NOME], para o fim especial de representá-la junto à SEMACE/Secretaria de Meio Ambiente de Iguatu.')
    
    doc.add_paragraph('\nIguatu-CE, {DATA_HOJE}.')
    
    doc.add_paragraph('\n\n____________________________________________________')
    doc.add_paragraph('{RAZAO_SOCIAL}')
    
    doc.save(os.path.join(TEMPLATES_DIR, 'Procuracao_SEMACE.docx'))

def create_requerimento():
    doc = Document()
    doc.add_heading('REQUERIMENTO PADRÃO', 0)
    
    doc.add_paragraph('Ilmo. Sr. Secretário de Meio Ambiente,')
    
    p = doc.add_paragraph('A empresa ')
    p.add_run('{RAZAO_SOCIAL}').bold = True
    p.add_run(' vem respeitosamente requerer a V.Sa. a emissão de Lincença Ambiental para sua atividade de {ATIVIDADE}.')
    
    doc.add_paragraph('Nestes Termos,\nPede Deferimento.')
    
    doc.add_paragraph('\nIguatu-CE, {DATA_HOJE}.')
    
    doc.save(os.path.join(TEMPLATES_DIR, 'Requerimento_Padrao.docx'))

if __name__ == "__main__":
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
    create_contrato()
    create_procuracao()
    create_requerimento()
    print("Templates criados com sucesso!")
